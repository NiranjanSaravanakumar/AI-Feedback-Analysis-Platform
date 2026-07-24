"""
utils/report.py
Generates a DOCX summary report from the enriched dataframe.
"""

import io
from typing import Dict, Any, List
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def generate_report(
    enriched_df: pd.DataFrame,
    cleaning_log: Dict[str, Any],
    sample_messages: Dict[str, List[Dict]],
    stats: Dict[str, Any],
) -> bytes:
    """
    Build a DOCX report and return it as bytes for Streamlit download.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Customer Feedback Intelligence Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, RGBColor(0x1e, 0x40, 0xaf))  # dark blue

    sub = doc.add_paragraph(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        f"System: QuickCart Feedback Intelligence System"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    _add_kv_table(doc, [
        ("Total Records Processed", str(cleaning_log.get("rows_loaded", "N/A"))),
        ("Records After Cleaning", str(cleaning_log.get("final_rows", "N/A"))),
        ("Rows Removed", str(
            cleaning_log.get("rows_loaded", 0) - cleaning_log.get("final_rows", 0)
        )),
        ("Report Date", datetime.datetime.now().strftime("%B %d, %Y")),
    ])

    # ── 2. Data Quality Summary ────────────────────────────────────────────────
    doc.add_heading("2. Data Quality Summary", level=1)
    _add_kv_table(doc, [
        ("Rows Loaded", str(cleaning_log.get("rows_loaded", "N/A"))),
        ("Duplicate Rows Removed", str(cleaning_log.get("duplicate_rows_removed", 0))),
        ("Blank Feedback Removed", str(cleaning_log.get("blank_feedback_removed", 0))),
        ("Timestamps Normalized", str(cleaning_log.get("invalid_timestamps_normalized", 0))),
        ("Timestamps Set Null (Unparseable)", str(cleaning_log.get("timestamps_set_null", 0))),
        ("Ratings Coerced to Null", str(cleaning_log.get("ratings_coerced", 0))),
        ("Source Values Normalized", str(cleaning_log.get("source_normalized", 0))),
        ("Final Rows", str(cleaning_log.get("final_rows", "N/A"))),
    ])

    # ── 3. Sentiment Breakdown ─────────────────────────────────────────────────
    doc.add_heading("3. Overall Sentiment Breakdown", level=1)
    total = stats.get("total", len(enriched_df))
    sentiment_rows = []
    for s in ("positive", "negative", "neutral"):
        count = stats.get(f"sentiment_{s}_count", 0)
        pct = stats.get(f"sentiment_{s}_pct", 0.0)
        sentiment_rows.append((s.capitalize(), str(count), f"{pct}%"))

    _add_header_table(
        doc,
        headers=["Sentiment", "Count", "Percentage"],
        rows=sentiment_rows,
    )

    # ── 4. Top 5 Complaint Categories ─────────────────────────────────────────
    doc.add_heading("4. Top 5 Complaint Categories", level=1)
    top_cats = enriched_df["category"].value_counts().head(5)
    cat_rows = []
    for i, (cat, count) in enumerate(top_cats.items(), 1):
        pct = round(count / total * 100, 1) if total else 0
        cat_rows.append((str(i), cat, str(count), f"{pct}%"))

    _add_header_table(
        doc,
        headers=["Rank", "Category", "Count", "% of Total"],
        rows=cat_rows,
    )

    # ── 5. Representative Examples per Category ────────────────────────────────
    doc.add_heading("5. Representative Feedback Examples", level=1)
    doc.add_paragraph(
        "The following examples were selected as representative entries for each top category, "
        "prioritizing negative sentiment as most actionable."
    ).runs[0].font.size = Pt(10)

    for cat, messages in sample_messages.items():
        doc.add_heading(f"  {cat}", level=2)
        for msg in messages:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{msg.get('sentiment', '?').upper()}] ").bold = True
            p.add_run(str(msg.get("feedback_text", "")))
            p.add_run(f"\n   → Summary: {msg.get('summary', '')}").italic = True
            p.add_run(
                f"\n   Source: {msg.get('source', 'N/A')} | "
                f"Rating: {msg.get('rating', 'N/A')} | "
                f"ID: {msg.get('id', 'N/A')}"
            ).font.size = Pt(9)
        doc.add_paragraph()

    # ── 6. AI Usage Log ───────────────────────────────────────────────────────
    doc.add_heading("6. AI Usage Notes", level=1)
    notes = [
        "NVIDIA LLaMA-3.1-8B-Instruct was used for sentiment, category classification, and summarization.",
        "Rows were processed in batches of 20 to optimize API usage and cost.",
        "All AI outputs were validated against a fixed allowed list (categories, sentiments).",
        "Any response outside the allowed values was replaced with 'Other' / 'neutral'.",
        "The AI was instructed to trust feedback text over star rating when they conflict.",
        "Failed batches were retried once; persistent failures received a neutral/Other fallback.",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet").runs[0].font.size = Pt(10)

    # ── Save to bytes ──────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_heading_color(heading, color: RGBColor):
    for run in heading.runs:
        run.font.color.rgb = color


def _add_kv_table(doc: Document, rows: List[tuple]):
    """Two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].text = val
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


def _add_header_table(doc: Document, headers: List[str], rows: List[tuple]):
    """Table with bold header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header
    header_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        cell = header_cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "1e40af")

    # Data rows
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = val
            cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()


def _shade_cell(cell, fill_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)
